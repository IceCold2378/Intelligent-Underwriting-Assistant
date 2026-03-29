"""
Document parsing service: extracts text from PDF, DOCX, and TXT files.
"""

import io
import logging

from PyPDF2 import PdfReader

from app.exceptions import DocumentParsingError, UnsupportedFileTypeError, FileTooLargeError
from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


def validate_file(filename: str, file_bytes: bytes) -> None:
    """Validate file type and size."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise UnsupportedFileTypeError(filename)

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > settings.MAX_FILE_SIZE_MB:
        raise FileTooLargeError(settings.MAX_FILE_SIZE_MB)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Extract text from a file based on its extension.
    Supports: PDF, DOCX, TXT
    """
    validate_file(filename, file_bytes)

    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext == "txt":
        return _extract_txt(file_bytes)
    else:
        raise UnsupportedFileTypeError(filename)


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise DocumentParsingError(
                "Could not extract any text from the PDF. "
                "The file may be scanned/image-based. OCR is not yet supported."
            )
        
        logger.info("Extracted %d characters from %d PDF pages", len(full_text), len(reader.pages))
        return full_text

    except DocumentParsingError:
        raise
    except Exception as e:
        logger.exception("PDF parsing failed: %s", e)
        raise DocumentParsingError(f"Failed to parse PDF: {str(e)}")


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise DocumentParsingError("Could not extract any text from the DOCX file.")
        
        logger.info("Extracted %d characters from DOCX", len(full_text))
        return full_text

    except DocumentParsingError:
        raise
    except ImportError:
        raise DocumentParsingError(
            "DOCX support requires the 'python-docx' package. "
            "Install it with: pip install python-docx"
        )
    except Exception as e:
        logger.exception("DOCX parsing failed: %s", e)
        raise DocumentParsingError(f"Failed to parse DOCX: {str(e)}")


def _extract_txt(file_bytes: bytes) -> str:
    """Extract text from plain text bytes."""
    try:
        text = file_bytes.decode("utf-8")
        if not text.strip():
            raise DocumentParsingError("The uploaded text file is empty.")
        logger.info("Extracted %d characters from TXT", len(text))
        return text
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
            return text
        except Exception as e:
            raise DocumentParsingError(f"Could not decode text file: {str(e)}")
